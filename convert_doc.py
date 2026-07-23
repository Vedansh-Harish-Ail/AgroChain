import re
import os
import time
import urllib.parse
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="1F4E79", sz="6", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def make_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def add_formatted_runs(paragraph, text):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 40, 40)
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 102, 204)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def build_docx(md_path, docx_path):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(34, 34, 34)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_table(lines):
        if not lines:
            return
        rows = []
        for line in lines:
            if '---' in line and not any(c.isalnum() for c in line.replace('-', '')):
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if cells and any(c for c in cells):
                rows.append(cells)
        if not rows:
            return
        
        num_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < num_cols:
                r.append('')
        
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        set_table_borders(table, color="1F4E79", sz="6")
        
        col_max_lens = [max(len(r[c]) for r in rows) for c in range(num_cols)]
        total_len = sum(col_max_lens) or 1
        col_widths = [max(1.0, (l / total_len) * 6.5) for l in col_max_lens]
        scale = 6.5 / sum(col_widths)
        col_widths = [w * scale for w in col_widths]

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            if r_idx == 0:
                make_row_header(row)
            else:
                prevent_row_split(row)

            for c_idx, cell_value in enumerate(row_data):
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    cell.width = Inches(col_widths[c_idx])
                    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.1
                    
                    if r_idx == 0:
                        set_cell_background(cell, "1F4E79")
                        add_formatted_runs(p, cell_value)
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10.5)
                    else:
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "F4F6F9")
                        else:
                            set_cell_background(cell, "FFFFFF")
                        add_formatted_runs(p, cell_value)
                        for run in p.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after = Pt(6)

    def flush_code(lines):
        if not lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        
        code_text = "".join(lines)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(40, 40, 40)
        
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        pPr.append(shd)

    i = 0
    while i < len(lines):
        line = lines[i]
        raw_line = line.rstrip('\r\n')
        stripped = raw_line.strip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                flush_code(code_lines)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table(table_lines)
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw_line + '\n')
            i += 1
            continue

        # Image Embeddings
        if stripped.startswith('![') and ']' in stripped and '(' in stripped and stripped.endswith(')'):
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False
            match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
            if match:
                caption = match.group(1)
                img_url = match.group(2)
                
                # Decode %20 and URL encoded characters
                unquoted_url = urllib.parse.unquote(img_url)
                clean_path = unquoted_url.replace('file:///', '').replace('file://', '').replace('/', '\\')
                clean_path = os.path.abspath(clean_path)
                
                if os.path.exists(clean_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run()
                    try:
                        run.add_picture(clean_path, width=Inches(5.8))
                    except Exception as err:
                        print(f"Failed to embed picture {clean_path}: {err}")
                    
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(12)
                    crun = cp.add_run(f"Figure: {caption}")
                    crun.italic = True
                    crun.font.name = 'Calibri'
                    crun.font.size = Pt(9.5)
                    crun.font.color.rgb = RGBColor(80, 80, 80)
                else:
                    print(f"Warning: Image file not found: {clean_path}")
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False

        if not stripped:
            i += 1
            continue

        # Horizontal Rule
        if stripped == '---' or stripped == '***':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_________________________________________________________________________________")
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            add_formatted_runs(p, stripped[2:])
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(31, 78, 121)
                run.font.size = Pt(20)
        elif stripped.startswith('## '):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_runs(p, stripped[3:])
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(46, 117, 182)
                run.font.size = Pt(15)
        elif stripped.startswith('### '):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, stripped[4:])
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(34, 34, 34)
                run.font.size = Pt(13)
        elif stripped.startswith('#### '):
            p = doc.add_heading(level=4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, stripped[5:])
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(80, 80, 80)
                run.font.size = Pt(11.5)
        # Lists
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^\d+\.\s', stripped)
            prefix_len = len(match.group(0))
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, stripped[prefix_len:])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_runs(p, stripped)

        i += 1

    if in_table:
        flush_table(table_lines)
    if in_code_block:
        flush_code(code_lines)

    saved = False
    for attempt in range(3):
        try:
            doc.save(docx_path)
            print(f"Successfully updated {docx_path} with embedded screenshots!")
            saved = True
            break
        except PermissionError:
            print(f"Attempt {attempt+1}: File {docx_path} is locked. Retrying in 1s...")
            time.sleep(1)
            
    if not saved:
        print(f"Error: {docx_path} is locked by Microsoft Word. Please close Word and run python script.")

if __name__ == '__main__':
    md_file = r"c:\MY PROJECTS\AgroChain-Morden\AgroChain_Documentation_Full.md"
    single_docx_file = r"c:\MY PROJECTS\AgroChain-Morden\AgroChain_Documentation_Full.docx"
    
    build_docx(md_file, single_docx_file)
